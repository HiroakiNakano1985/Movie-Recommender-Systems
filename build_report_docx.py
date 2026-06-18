# -*- coding: utf-8 -*-
"""Build the final report as an editable Word document (REPORT.docx).

Mirrors REPORT.md but with native Word headings, tables and embedded figures so
it can be tweaked in Word and exported to PDF in one click.

Run (after `python main.py` has produced results/figures):
    python build_report_docx.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "results" / "figures"
OUT = ROOT / "REPORT.docx"

ACCENT = RGBColor(0x2E, 0x6F, 0xF2)
GREY = RGBColor(0x55, 0x5E, 0x6E)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def para(text="", italic=False, color=None, size=None, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def figure(name, caption, width=6.0):
    path = FIG / name
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = para(caption, italic=True, color=GREY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        return cap


def table(data, header=True):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(data):
        for c, val in enumerate(row):
            cell = t.cell(r, c)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    if header and r == 0:
                        run.bold = True
    return t


# ===== Title block =====
title = doc.add_heading("Movie Recommender Systems — Individual Project Report", level=0)
para("Author: Hiroaki Nakano", bold=True)
para("Course: Recommender Systems — Prof. Marc Torrens, Esade")
para("Date: June 2026   |   Track: Movies — MovieLens Latest Small")
p = para("Live demo: ", bold=True)
p.add_run("https://movie-recommender-systems-bgvfxcjsf7myfvfxjyircb.streamlit.app").font.color.rgb = ACCENT
p2 = para("Code: ", bold=True)
p2.add_run("https://github.com/HiroakiNakano1985/Movie-Recommender-Systems").font.color.rgb = ACCENT

# ===== 1. Introduction =====
doc.add_heading("1. Introduction", level=1)
para(
    "The goal of this project is to build a single movie-recommendation prototype "
    "and grow it method by method, so that the same data and the same evaluation "
    "protocol can be used to compare a family of algorithms fairly. Rather than "
    "chasing a single “best” model, the project treats recommendation as a "
    "multi-objective problem: accuracy matters, but so do catalog coverage, novelty, "
    "diversity and popularity bias. The deliverable is therefore both a set of "
    "working, comparable algorithms and a critical analysis of when and why each one "
    "is appropriate."
)
para(
    "Concretely, the prototype implements seven recommenders — three "
    "non-personalized baselines, a content-based model, item-item and user-user "
    "collaborative filtering, and a matrix-factorization model — evaluates them "
    "with a common top-N protocol plus beyond-accuracy and rating-prediction "
    "metrics, and exposes everything through an interactive Streamlit application."
)
para(
    "All code is organized as a small Python package (src/), orchestrated by "
    "main.py, with two analysis notebooks (notebooks/01_eda.ipynb, "
    "notebooks/02_comparison.ipynb) and the demo app (app.py)."
)

# ===== 2. Dataset =====
doc.add_heading("2. Dataset description", level=1)
para(
    "I use the MovieLens Latest Small dataset from GroupLens "
    "(https://grouplens.org/datasets/movielens/), free for research and educational "
    "use (see data/raw/ml-latest-small/README.txt). It consists of ratings.csv "
    "(userId, movieId, rating, timestamp) and movies.csv (movieId, title, genres)."
)
table([
    ["Property", "Value"],
    ["Users", "610"],
    ["Movies (rated)", "9,724"],
    ["Movies (catalog)", "9,742"],
    ["Ratings", "100,836"],
    ["Rating scale", "0.5 – 5.0 (mean 3.5, mode 4.0)"],
    ["Density", "1.7% (sparsity 98.3%)"],
    ["Ratings per user", "min 20, median 70, max 2,698"],
    ["Ratings per item", "min 1, median 3, max 329"],
])
para(
    "Two facts drive every design decision. First, the matrix is 98.3% sparse — "
    "which is normal (even relatively dense) by recommender standards. Second, the "
    "sparsity is asymmetric: users are well-described (median 70 ratings) while items "
    "are thin (median 3). The weak axis is the item axis, and that is exactly where "
    "collaborative methods struggle."
)

# ===== 3. Preprocessing & EDA =====
doc.add_heading("3. Preprocessing and EDA", level=1)
para(
    "Loading and validation live in src/data_loading.py. The EDA is in "
    "notebooks/01_eda.ipynb; the key figures are reproduced below."
)
para("Rating distribution.", bold=True)
para(
    "Ratings skew positive (mean 3.5, mode 4.0). This positivity bias is why I "
    "define relevance as a rating ≥ 4.0 in the evaluation: simply being rated is "
    "not the same as being liked."
)
figure("rating_distribution.png", "Figure 1. Rating distribution (positivity bias).", 5.0)
para("Activity distributions and the long tail.", bold=True)
para(
    "A few power users and a small head of blockbuster movies dominate. About 20% of "
    "the most popular movies account for 80% of all ratings — the classic long "
    "tail, and the reason a most-popular baseline is hard to beat on raw accuracy."
)
figure("long_tail.png", "Figure 2. The long tail of item popularity.", 5.0)
para("Train/test split.", bold=True)
para(
    "train_test_split_ratings performs a per-user stratified 80/20 split: the unit "
    "of splitting is the individual rating, but stratified by user so that every user "
    "keeps ~80% of their ratings for training and ~20% as held-out test items. This "
    "is required for top-N evaluation — we must build a profile for each user "
    "from train and check held-out positives from test. All 610 users appear in both "
    "splits (80,668 train / 20,168 test)."
)

# ===== 4. Algorithms =====
doc.add_heading("4. Algorithms implemented", level=1)
para(
    "All recommenders share the same interface: fit(...) and "
    "recommend(user_id, ratings_train, n, exclude_seen) returning (item_id, score) "
    "pairs. This uniformity is what makes a fair comparison possible."
)
bullet("Most Popular (rank by # ratings), Highest Average (mean rating, min 20 ratings), Random (lower bound / coverage reference).", "Non-personalized baselines: ")
bullet("each movie is a TF-IDF vector over genres; a user profile is the mean-centered, rating-weighted sum of their rated movies’ vectors; recommend the unseen items most cosine-similar to that profile. Immune to user-item sparsity.", "Content-based: ")
bullet("item-item and user-user k-NN with adjusted cosine. Sparsity safeguards: min_support (drop items with < 5 ratings) and shrinkage sim_adj = n_ij/(n_ij+λ)·sim. For top-N ranking the models rank by the unnormalized weighted sum; normalization is kept only for rating prediction.", "Collaborative filtering: ")
bullet("biased model r̂(u,i) = μ + b_u + b_i + p_u·q_i trained with SGD on observed ratings (regularized squared error). Implemented from scratch in NumPy (scikit-surprise does not build on Python 3.14), which also keeps it fully transparent. The canonical answer to sparsity: dense latent vectors learned from observed ratings.", "Matrix factorization: ")

# ===== 5. Protocol =====
doc.add_heading("5. Evaluation protocol", level=1)
bullet("a test item is relevant for a user if rated ≥ 4.0.", "Relevance: ")
bullet("Precision@K, Recall@K, NDCG@K, MRR@K, Hit-Rate@K.", "Top-N ranking (K=10): ")
bullet("catalog coverage, novelty (mean −log₂ P(item)), intra-list diversity (1 − cosine over genre vectors).", "Beyond-accuracy: ")
bullet("RMSE and MAE for models exposing predict_score.", "Rating prediction: ")
para(
    "Recommendations are generated from training data only, excluding items already "
    "seen in training. Metrics are averaged over the 599 users with at least one "
    "relevant test item.", italic=True, color=GREY
)

# ===== 6. Results =====
doc.add_heading("6. Results", level=1)
para("6.1 Ranking and beyond-accuracy (K = 10)", bold=True)
table([
    ["Method", "NDCG", "P@10", "Recall", "MRR", "HitRate", "Coverage", "Novelty", "Diversity"],
    ["User-User CF", "0.219", "0.163", "0.145", "0.393", "0.648", "0.031", "2.19", "0.75"],
    ["Item-Item CF", "0.211", "0.155", "0.126", "0.402", "0.644", "0.086", "2.80", "0.72"],
    ["Most Popular", "0.161", "0.122", "0.104", "0.312", "0.574", "0.006", "1.66", "0.76"],
    ["Matrix Factorization", "0.072", "0.056", "0.041", "0.168", "0.356", "0.049", "3.41", "0.78"],
    ["Highest Average", "0.062", "0.046", "0.034", "0.151", "0.301", "0.004", "3.39", "0.66"],
    ["Content-Based", "0.008", "0.007", "0.005", "0.019", "0.063", "0.341", "7.28", "0.085"],
    ["Random", "0.001", "0.001", "0.001", "0.004", "0.013", "0.489", "7.57", "0.80"],
])
para("Sorted by NDCG@10 (rank-aware accuracy). Source: results/metrics.csv.", italic=True, color=GREY, size=9)
figure("accuracy_comparison.png", "Figure 3. Ranking accuracy by method.", 5.5)

para("6.2 The accuracy vs. coverage trade-off", bold=True)
para(
    "The single most important result is the frontier: the most accurate methods "
    "(CF) recommend from a tiny slice of the catalog, while high-coverage methods "
    "(content-based, random) are inaccurate. No method is simultaneously accurate and "
    "high-coverage; a real product must choose a point on this frontier."
)
figure("accuracy_vs_coverage.png", "Figure 4. Accuracy vs. coverage trade-off.", 5.0)

para("6.3 Novelty and diversity", bold=True)
para(
    "Novelty rises as methods move away from the popular head. The striking case is "
    "the content-based model: highest novelty but lowest intra-list diversity (0.085) "
    "— it keeps recommending items from the user’s single favourite genre "
    "cluster, varied across the catalog (novel) but monotonous within one list."
)
figure("novelty_diversity.png", "Figure 5. Novelty and intra-list diversity.", 6.0)

para("6.4 Rating prediction", bold=True)
table([
    ["Method", "RMSE", "MAE"],
    ["Matrix Factorization", "0.885", "0.679"],
    ["Item-Item CF", "0.905", "0.669"],
])
para("Global-mean baseline RMSE = 1.039; item-mean baseline RMSE = 0.976. Source: results/rating_metrics.csv.", italic=True, color=GREY, size=9)
figure("rating_prediction.png", "Figure 6. Rating-prediction error (lower is better).", 5.0)
para(
    "Matrix factorization is the weakest top-N model but the best rating predictor, "
    "improving on the global-mean baseline by ~15%. This is the key nuance: MF "
    "minimizes squared error, which is not a ranking objective, so a low RMSE does "
    "not translate into a good top-10. “MF doesn’t work” would be the "
    "wrong conclusion — it optimizes a different objective."
)

para("6.5 Ablation: TF-IDF vs. raw genre vectors", bold=True)
para(
    "The content-based model vectorizes each movie over its genres (one movie = one "
    "document, each genre = one token). I compared TF-IDF weighting against plain raw "
    "genre vectors (L2-normalized 0/1 presence, every genre equal). Since each genre "
    "appears at most once per movie the term-frequency is always 1, so the only thing "
    "TF-IDF changes is the IDF term — down-weighting ubiquitous genres (Drama, "
    "Comedy) and up-weighting rare, distinctive ones (Film-Noir, Western, IMAX)."
)
table([
    ["Variant", "NDCG", "P@10", "Recall", "HitRate", "Coverage", "Novelty", "Diversity"],
    ["TF-IDF genres", "0.0076", "0.0068", "0.0047", "0.063", "0.341", "7.28", "0.085"],
    ["Raw genre vectors", "0.0056", "0.0058", "0.0035", "0.053", "0.352", "7.29", "0.075"],
])
para("Same metric set as the main results table. Source: results/ablation_tfidf.csv; diversity in a fixed shared genre space.", italic=True, color=GREY, size=9)
para(
    "TF-IDF wins 5 of the 7 metrics — all four accuracy metrics (NDCG +36% "
    "relative, 0.0056 → 0.0076; hit-rate 0.053 → 0.063) plus intra-list diversity. By "
    "emphasizing distinctive genres it builds more discriminative user profiles. Raw "
    "genre vectors only edge ahead on coverage and novelty, because weighting every "
    "genre equally spreads recommendations slightly wider across the catalog. The "
    "overall gain is modest because the genre vocabulary is tiny (~20 tokens) and TF "
    "is always 1, so only the IDF term is active; a richer text signal (tags, "
    "synopsis) would give TF-IDF far more to work with."
)

# ===== 7. Examples =====
doc.add_heading("7. Recommendation examples", level=1)
para("Top recommendations for two contrasting users (from main.py):")
para("User 1 — 186 ratings, avg 4.41 (an enthusiast):", bold=True)
bullet("Pulp Fiction · Shawshank Redemption · Braveheart · Terminator 2", "Most Popular: ")
bullet("Pulp Fiction · Snatch · Memento · The Godfather · Back to the Future", "Item-Item CF: ")
bullet("Anastasia · Tarzan · Up · Dumbo · Free Willy (animation cluster)", "Content-Based: ")
para("User 599 — 1,982 ratings, avg 2.64 (a critical heavy user):", bold=True)
bullet("The Matrix · Star Wars IV · Schindler’s List · Empire Strikes Back", "Most Popular: ")
bullet("Star Wars IV · The Matrix · Empire Strikes Back · LotR: Fellowship", "User-User CF: ")
bullet("Lawrence of Arabia · Donnie Darko · Kill Bill 2 · Some Like It Hot", "Matrix Factorization: ")
para(
    "The qualitative differences match the metrics: CF leans on well-loved classics; "
    "content-based produces tight, same-genre lists; MF mixes acclaimed films from "
    "the learned latent space. These can be explored interactively, with "
    "per-recommendation explanations, in the Streamlit app."
)

# ===== 8. Limitations =====
doc.add_heading("8. Discussion of limitations", level=1)
bullet("results come from one 80/20 split; cross-validation or a temporal (leave-last-out by timestamp) split would be more robust and realistic.", "Single random split: ")
bullet("relevance is binary at ≥ 4.0; results shift with the threshold and ignore graded relevance.", "Fixed threshold: ")
bullet("popular items appear often in the test set, so offline accuracy rewards popularity; online A/B testing would judge discovery differently.", "Popularity dominates offline top-N: ")
bullet("an RMSE objective handicaps MF on ranking; a pairwise/implicit objective (BPR, ALS, logistic MF) would be a fairer latent-factor competitor.", "MF objective: ")
bullet("diversity is measured over genres only; richer features (tags, cast, synopsis embeddings) would sharpen both the model and the metric.", "Content signal is genres only: ")
bullet("new users/items are not analysed separately, although that is where content-based and MF would shine.", "Cold start not isolated: ")

# ===== 9. Conclusion =====
doc.add_heading("9. Conclusion", level=1)
para(
    "Building one prototype and comparing seven methods on identical data made the "
    "trade-offs concrete. Neighbourhood collaborative filtering — user-user "
    "slightly ahead of item-item, thanks to the denser user axis — gives the "
    "best top-N ranking. Most-popular is a deceptively strong accuracy baseline but "
    "reaches only 0.6% of the catalog, a textbook case of popularity bias. Matrix "
    "factorization is the best rating predictor yet a weak ranker, a clean "
    "illustration that the optimization objective must match the task. And the "
    "accuracy–coverage–novelty trade-offs show that “best” is "
    "product-dependent, not absolute."
)
para(
    "The most valuable engineering lessons were about sparsity: standard methods do "
    "work at 98% sparsity, provided each method is matched to the data’s weak "
    "axis (MF and content-based for the thin item tail) and neighbourhood CF is "
    "protected with shrinkage, minimum support, and — crucially — ranked by "
    "the right score. Natural next steps are a hybrid CF + content model to cover "
    "cold/long-tail items and a ranking-objective matrix factorization to make latent "
    "factors competitive on top-N."
)
para(
    "A final, broader caveat concerns evaluation itself. Moving “beyond accuracy” to "
    "coverage, novelty and diversity — and, in production, to long-term retention or "
    "satisfaction — does not escape bias; it relocates it. There is no unbiased "
    "objective: every metric encodes both a value judgement (whose interest is being "
    "optimised) and a measurement process, each with its own bias. Offline accuracy "
    "inherits the exposure / feedback bias of the policy that generated the data; "
    "retention suffers survivorship bias (churned users leave no trace) and can reward "
    "compulsive rather than genuinely satisfying engagement; and by Goodhart’s law any "
    "single target, once optimised hard, decouples from the true — and unobservable — "
    "goal of long-term user utility. The practical response is not to hunt for the one "
    "“correct” metric but to make the biases explicit and manage them: optimise a "
    "basket of metrics with guardrails, inject exploration and use counterfactual "
    "estimators to debias feedback loops, analyse results per user segment, and keep "
    "long-term holdouts and human judgement in the loop. This is the deeper reading of "
    "“accuracy is not enough” — its replacements are not neutral either."
)

# ===== Appendix =====
doc.add_heading("Appendix — How to reproduce", level=1)
pd_ = para("Live demo (no setup needed): ", bold=True)
pd_.add_run("https://movie-recommender-systems-bgvfxcjsf7myfvfxjyircb.streamlit.app").font.color.rgb = ACCENT
code = doc.add_paragraph()
code.add_run(
    "pip install -r requirements.txt\n"
    "python main.py                 # full pipeline -> results/metrics.csv\n"
    "python build_artifacts.py      # cache trained models for the app\n"
    "streamlit run app.py           # interactive prototype\n"
    "python build_slides.py         # rebuild the slide deck"
).font.name = "Consolas"
for r in code.runs:
    r.font.size = Pt(9.5)

doc.save(OUT)
print(f"Saved report to {OUT} ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")
